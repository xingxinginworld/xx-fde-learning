"""md2docx.py — 轻量 Markdown -> Word 转换器（针对本项目交付文档定制）

支持的语法：
  #/##/###/#### 标题、> 引用、- 列表、- [ ] 勾选、1. 有序列表、
  | 表格 |、``` 代码块、--- 分隔线、**加粗**、行内 `code`

用法：
  python md2docx.py 输入.md [输出.docx]
  不带输出名时，自动在同目录生成同名 .docx

依赖：python-docx（全局已装 1.2.0）
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # 标题深蓝


def set_cjk(run, font="Microsoft YaHei"):
    """给 run 设置中英文字体（保证中文不乱码）。"""
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def shade(paragraph, fill="F2F4F7"):
    """给段落加浅灰底纹（用于代码块）。"""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def hr(paragraph):
    """给段落底部加一条横线（模拟 ---）。"""
    ppr = paragraph._p.get_or_add_pPr()
    pbd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pbd.append(bottom)
    ppr.append(pbd)


def add_inline(paragraph, text):
    """处理行内 **加粗** 与 `code`。"""
    # 先按 ** 粗体切，再在每段里处理 `
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_cjk(run)
        else:
            # 行内 code
            subs = re.split(r"(`[^`]+`)", part)
            for s in subs:
                if s.startswith("`") and s.endswith("`"):
                    r = paragraph.add_run(s[1:-1])
                    r.font.name = "Consolas"
                    r._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
                    r._element.rPr.find(qn("w:rFonts")).set(qn("w:eastAsia"), "Consolas")
                else:
                    r = paragraph.add_run(s)
                    set_cjk(r)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
        run._element.rPr.find(qn("w:rFonts")).set(qn("w:eastAsia"), "Consolas")
        shade(p)
    doc.add_paragraph()


def add_table(doc, rows):
    # rows[0] 表头，rows[1] 分隔行(丢弃)，其余数据
    header = rows[0]
    sep = "".join(rows[1]) if len(rows) > 1 else ""
    data = rows[2:] if len(rows) > 2 and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", sep) else rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h.strip())
        run.bold = True
        set_cjk(run)
    # 数据
    for r in data:
        cells = table.add_row().cells
        for i, c in enumerate(r):
            cells[i].text = ""
            add_inline(cells[i].paragraphs[0], c.strip())
    doc.add_paragraph()


def convert(md_text):
    doc = Document()
    # 默认正文字体（中文）
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Microsoft YaHei"
    nrpr = normal.element.get_or_add_rPr()
    nf = nrpr.find(qn("w:rFonts"))
    if nf is None:
        nf = OxmlElement("w:rFonts")
        nrpr.append(nf)
    nf.set(qn("w:eastAsia"), "Microsoft YaHei")
    nf.set(qn("w:ascii"), "Microsoft YaHei")
    nf.set(qn("w:hAnsi"), "Microsoft YaHei")

    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # 1) 代码块
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            # mermaid 等无语言标注也当代码块渲染
            add_code_block(doc, buf)
            continue

        # 2) 分隔线
        if re.match(r"^\s*---+\s*$", line) and (i == 0 or lines[i - 1].strip() == ""):
            p = doc.add_paragraph()
            hr(p)
            i += 1
            continue

        # 3) 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(level=level)
            add_inline(p, m.group(2).strip())
            # 标题配色
            for run in p.runs:
                run.font.color.rgb = ACCENT
            i += 1
            continue

        # 4) 引用
        if line.startswith("> "):
            buf = []
            while i < n and lines[i].startswith("> "):
                buf.append(lines[i][2:])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, " ".join(buf).strip())
            shade(p, "EAF1F8")
            i += 0
            continue

        # 5) 表格
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                row = [c for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            add_table(doc, rows)
            continue

        # 6) 列表（有序 / 无序 / 勾选）
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        m_ul = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m_ol.group(3))
            i += 1
            continue
        if m_ul:
            content = m_ul.group(2)
            if re.match(r"^\[[ xX]\]\s+", content):
                content = re.sub(r"^\[[ xX]\]\s+", "☐ ", content)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue

        # 7) 空行
        if line.strip() == "":
            i += 1
            continue

        # 8) 普通段落
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1

    return doc


def main():
    if len(sys.argv) < 2:
        print("用法: python md2docx.py 输入.md [输出.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    text = src.read_text(encoding="utf-8")
    doc = convert(text)
    doc.save(str(out))
    print(f"✅ {src.name} -> {out.name}")


if __name__ == "__main__":
    main()
